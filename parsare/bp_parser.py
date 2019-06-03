# https://python-docx.readthedocs.io/en/latest/user/quickstart.html#
# https://python-docx.readthedocs.io/en/latest/user/documents.html


# process name, subsheet, collecction

import sys,os
import xml.etree.ElementTree as ET
from docx import Document

##try:
##    from docx import Document
##except ModuleNotFoundError:
##    try:
##        os.system('pip install python-docx')
##    except:
##        print("could not install it")


        
    

def check_element(e):
    """
        recursive function for getting all of the of node children
        """
    children = len(list(e))
    #print(f"Element {e.text} has {children} children")
    if children:
        for i in e:
            check_element(i)
            check_attr(e,'Collection Example')
    else:
        #print("Tag: {0} and Text: {1} \n".format(e.tag, e.text))
        check_attr(e,'Collection Example')

        
def check_attr(e,attr_name):
    """ e: xml.etree.ElementTree.Element
        s: string for value in attrib dictionary"""
    try:
        # get keys that have the value == attr_name where e.attrib is a dict
        attr_list = [k for k,v in e.attrib.items() if v == attr_name]
        for i in attr_list:
            #print( e.get(i))
            print(f"Element tag {e.tag} : with text {e.text} \t")

            #write to doc
            createWord("text.docx",str(e.attrib) )
    except:
        pass
        
    

def parseit(f = None):
    if f:
        # create root instance with tag and attributes as a dict
        tree = ET.parse(f)
        root = tree.getroot()
        
        # parse all node from root to and call chec_attr("searchword") on each node
        for c in root:
            check_element(c)
        

def createWord(name, p = "Add text for paragrph"):
    """ Create or open word document with name.docx
    """

    if name.endswith(".docx"):
        # verify if doc is present on disk
        if name in os.listdir(os.getcwd()):
            document = Document(name)
        else:    
            document = Document()
    else:
        print("Add extension in doc name e.g: name.docx")

    #prelucrare
    document.add_paragraph(p)
    #save the doc    
    document.save(name)

    
        
if __name__ == "__main__":

    parseit('bp')

    
##    r = ET.parse('bp').getroot()
   
    
